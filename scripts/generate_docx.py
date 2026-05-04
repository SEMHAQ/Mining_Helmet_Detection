"""将论文全文.md转换为DOCX格式，严格遵循《矿业研究与开发》投稿模板

模板要求：
- 一级标题：小四宋体加粗
- 二级标题：5号宋体加粗
- 正文：5号宋体，1.5倍行距，首行缩进2字符
- 参考文献：6号宋体
- 摘要：楷体小五
- 图表标题中英双语
- 颜色图需标注"颜色标识见电子版"
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── 字体/段落辅助 ──────────────────────────────────────────────

def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(10.5)):
    run.font.name = en_font
    run.font.size = size
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)


def set_para_fmt(p, line_spacing=1.5, first_indent=None,
                 space_before=Pt(0), space_after=Pt(0)):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_indent is not None:
        pf.first_line_indent = first_indent


# ── 文档元素 ───────────────────────────────────────────────────

def add_title(doc, text):
    """论文标题：黑体 小二(18pt) 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, cn_font='黑体', size=Pt(18))
    set_para_fmt(p, space_after=Pt(6))


def add_center_text(doc, text, cn_font='宋体', size=Pt(10.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size)
    set_para_fmt(p)


def add_label(doc, text, cn_font='黑体', size=Pt(10.5)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size)
    run.bold = True
    set_para_fmt(p)
    return p


def add_abstract_text(doc, text, cn_font='楷体', size=Pt(9)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size)
    set_para_fmt(p)


def add_keywords(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    set_run_font(r1, cn_font='黑体', size=Pt(10.5))
    r1.bold = True
    r2 = p.add_run(text)
    set_run_font(r2, size=Pt(10.5))
    set_para_fmt(p)


def add_heading(doc, text, level):
    """一级标题：小四(12pt)宋体加粗；二级标题：5号(10.5pt)宋体加粗"""
    if level == 1:
        size = Pt(12)   # 小四
        cn_font = '宋体'
        sb = Pt(12)
    elif level == 2:
        size = Pt(10.5)  # 五号
        cn_font = '宋体'
        sb = Pt(6)
    else:
        size = Pt(10.5)
        cn_font = '宋体'
        sb = Pt(3)
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size)
    run.bold = True
    set_para_fmt(p, space_before=sb, space_after=Pt(6))


def add_body(doc, text):
    """正文：5号宋体(10.5pt)，首行缩进2字符"""
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_run_font(run, size=Pt(10.5))
        if i % 2 == 1:
            run.bold = True
    set_para_fmt(p, first_indent=Cm(0.74))


def add_formula_ph(doc, tag):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'（公式 {tag}）')
    set_run_font(run, size=Pt(10.5))
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    set_para_fmt(p, space_before=Pt(6), space_after=Pt(6))


def add_image_ph(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[{caption}]')
    set_run_font(run, size=Pt(10.5))
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    set_para_fmt(p, space_before=Pt(6), space_after=Pt(6))


def add_color_note(doc):
    """颜色标识说明"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('（颜色标识见电子版）')
    set_run_font(run, size=Pt(9))
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    set_para_fmt(p, space_after=Pt(3))


def add_table_caption(doc, cn, en):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cn)
    set_run_font(r, cn_font='黑体', size=Pt(9))
    r.bold = True
    set_para_fmt(p, space_before=Pt(6), space_after=Pt(3))
    if en:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(en)
        set_run_font(r2, size=Pt(9))
        r2.bold = True
        set_para_fmt(p2, space_after=Pt(3))


def add_fig_caption(doc, cn, en):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cn)
    set_run_font(r, size=Pt(9))
    r.bold = True
    set_para_fmt(p, space_after=Pt(3))
    if en:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(en)
        set_run_font(r2, size=Pt(9))
        r2.bold = True
        set_para_fmt(p2, space_after=Pt(6))


def add_three_line_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h.strip())
        set_run_font(run, size=Pt(9))
        run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val.strip())
            set_run_font(run, size=Pt(9))
    _set_three_line_borders(table)


def _set_three_line_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = tbl.makeelement(qn('w:tblPr'), {})
        tbl.insert(0, tblPr)
    for b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(b)
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ['top', 'bottom']:
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '12',
            qn('w:space'): '0', qn('w:color'): '000000'})
        borders.append(el)
    insideH = borders.makeelement(qn('w:insideH'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:space'): '0', qn('w:color'): '000000'})
    borders.append(insideH)
    for edge in ['left', 'right', 'insideV']:
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): 'auto'})
        borders.append(el)
    tblPr.append(borders)
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = tc.makeelement(qn('w:tcPr'), {})
                tc.insert(0, tcPr)
            tcB = tcPr.makeelement(qn('w:tcBorders'), {})
            bot = tcB.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single', qn('w:sz'): '6',
                qn('w:space'): '0', qn('w:color'): '000000'})
            tcB.append(bot)
            tcPr.append(tcB)


def add_ref(doc, text):
    """参考文献：6号(7.5pt)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=Pt(7.5))
    set_para_fmt(p, space_after=Pt(2))


def add_footnote(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=Pt(9))
    set_para_fmt(p, first_indent=Cm(0))


# ── 主转换 ─────────────────────────────────────────────────────

# 需要标注"颜色标识见电子版"的图（有颜色区分的对比图）
COLOR_FIGURES = {'图5', '图6', '图7', '图8', '图9', '图10',
                 '图11', '图12', '图13', '图14'}


def generate_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    blocks = _split_by_h2(content)

    # ── 1. 中文标题 ──
    add_title(doc, '基于ECB-YOLO的井下安全头盔检测模型')

    # ── 2. 作者 ──
    add_center_text(doc, '彭东海¹，余焕杰²，江华晋²')

    # ── 3. 单位 ──
    add_center_text(doc, '¹ 韶关学院 信息工程学院，广东 韶关 512005，中国', size=Pt(9))
    add_center_text(doc, '² 湖南工商大学 计算机学院，湖南 长沙 410205，中国', size=Pt(9))

    # ── 4-6. 中文摘要区块 ──
    abs_block = blocks.get('摘要', '')
    _render_abstract_block(doc, abs_block)

    # ── 7. 英文标题 ──
    add_title(doc, 'Underground Safety Helmet Detection Model Based on ECB-YOLO')

    # ── 8. 英文作者 ──
    add_center_text(doc, 'PENG Donghai¹, YU Huanjie², JIANG Huajin²')
    add_center_text(doc,
        '¹ School of Information Engineering, Shaoguan University, '
        'Shaoguan, Guangdong 512005, China', size=Pt(9))
    add_center_text(doc,
        '² School of Computer Science, Hunan University of Technology and Business, '
        'Changsha, Hunan 410205, China', size=Pt(9))

    # ── 9-10. 英文摘要 ──
    en_block = blocks.get('Underground Safety Helmet Detection Model Based on ECB-YOLO', '')
    _render_en_abstract(doc, en_block)

    # ── 正文 ──
    body_keys = [k for k in blocks if re.match(r'^\d+\s', k)]
    for key in body_keys:
        add_heading(doc, key, level=1)
        _render_body_block(doc, blocks[key])

    # ── 参考文献 ──
    ref_block = blocks.get('参考文献（References）', blocks.get('参考文献', ''))
    if ref_block:
        add_heading(doc, '参考文献', level=1)
        _render_refs(doc, ref_block)

    doc.save(docx_path)
    print(f'DOCX 已生成: {docx_path}')


def _split_by_h2(content):
    parts = re.split(r'^## (.+)$', content, flags=re.MULTILINE)
    blocks = {}
    i = 1
    while i < len(parts) - 1:
        title = parts[i].strip()
        body = parts[i + 1] if i + 1 < len(parts) else ''
        blocks[title] = body
        i += 2
    return blocks


def _strip_bold(text):
    return re.sub(r'\*\*(.*?)\*\*', r'\1', text)


def _render_abstract_block(doc, block):
    lines = [l.strip() for l in block.split('\n') if l.strip()]
    abstract_done = False
    for line in lines:
        if line == '---':
            continue
        if not abstract_done and not line.startswith('**') and not line.startswith('收稿'):
            add_label(doc, '摘要')
            add_abstract_text(doc, line)
            abstract_done = True
            continue
        if '关键词' in line and '**' in line:
            m = re.search(r'：\s*(.*)', _strip_bold(line))
            if m:
                add_keywords(doc, '关键词：', m.group(1).strip())
            continue
        if '中图分类号' in line:
            add_body(doc, _strip_bold(line))
            continue
        if any(line.startswith(k) for k in ('收稿日期', '基金项目', '作者简介', '通信作者')):
            add_footnote(doc, line)
            continue
        if line.startswith('**'):
            add_footnote(doc, _strip_bold(line))


def _render_en_abstract(doc, block):
    lines = [l.strip() for l in block.split('\n') if l.strip()]
    abstract_printed = False
    for line in lines:
        if line == '---':
            continue
        if line.startswith('**Keywords'):
            m = re.search(r'Keywords[：:]\s*(.*)', _strip_bold(line))
            if m:
                add_keywords(doc, 'Keywords: ', m.group(1).strip())
            continue
        if 'China' in line and len(line) < 120:
            continue
        if line.startswith('PENG ') or line.startswith('YU ') or line.startswith('JIANG '):
            continue
        if not abstract_printed and len(line) > 100:
            add_label(doc, 'Abstract', cn_font='Times New Roman')
            add_abstract_text(doc, line, cn_font='Times New Roman')
            abstract_printed = True


def _render_body_block(doc, block):
    lines = block.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line or line == '---':
            i += 1
            continue

        # 三级标题
        if line.startswith('### ') and not line.startswith('#### '):
            add_heading(doc, line[4:].strip(), level=2)
            i += 1
            continue

        # 四级标题
        if line.startswith('#### '):
            add_heading(doc, line[5:].strip(), level=3)
            i += 1
            continue

        # 图片 → 占位符
        if line.startswith('!['):
            m = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if m:
                caption = m.group(1)
                add_image_ph(doc, caption)
                # 对比图/PR曲线/训练曲线/混淆矩阵 → 标注颜色见电子版
                for fig_id in COLOR_FIGURES:
                    if fig_id in caption:
                        add_color_note(doc)
                        break
            i += 1
            continue

        # 图片标题
        if line.startswith('**图') and line.endswith('**'):
            cn = _strip_bold(line)
            en = ''
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('**Fig.'):
                i += 1
                en = _strip_bold(lines[i].strip())
            add_fig_caption(doc, cn, en)
            i += 1
            continue

        # 表格标题
        if line.startswith('**表') and line.endswith('**'):
            cn = _strip_bold(line)
            en = ''
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('**Table'):
                i += 1
                en = _strip_bold(lines[i].strip())
            add_table_caption(doc, cn, en)
            i += 1
            continue

        # 跳过单独的 Fig./Table 行
        if line.startswith('**Fig.') or line.startswith('**Table'):
            i += 1
            continue

        # LaTeX 公式 → 占位符
        if line.startswith('$$'):
            first = line[2:]
            if first.endswith('$$'):
                # 单行公式：$$...\tag{N}$$
                formula_text = first[:-2].strip()
            else:
                # 多行公式
                formula_lines = [first]
                i += 1
                while i < len(lines) and not lines[i].strip().endswith('$$'):
                    formula_lines.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    formula_lines.append(lines[i].strip()[:-2])
                formula_text = ' '.join(formula_lines).strip()
            tag_m = re.search(r'\\tag\{(\d+)\}', formula_text)
            tag = tag_m.group(1) if tag_m else '?'
            add_formula_ph(doc, tag)
            i += 1
            continue

        # 表格
        if '|' in line and not line.startswith('**'):
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i] and not lines[i].strip().startswith('**'):
                table_lines.append(lines[i].strip())
                i += 1
            headers = []
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|') if c.strip()]
                if all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                    continue
                if not headers:
                    headers = cells
                else:
                    rows.append(cells)
            if headers and rows:
                add_three_line_table(doc, headers, rows)
            continue

        # 注释行
        if line.startswith('注：') or line.startswith('注¹') or line.startswith('注²'):
            add_footnote(doc, line)
            i += 1
            continue

        # 加粗独立行
        if line.startswith('**') and '**' in line[2:]:
            add_body(doc, _strip_bold(line))
            i += 1
            continue

        # 普通正文（合并连续行，遇到空行/特殊标记停止）
        para_lines = [line]
        i += 1
        while i < len(lines):
            ns = lines[i].strip()
            if not ns:
                break
            if ns.startswith('#') or ns.startswith('![') or ns.startswith('$$') or ns == '---':
                break
            if '|' in ns and not ns.startswith('**'):
                break
            if ns.startswith('**') and ('图' in ns or '表' in ns or 'Fig' in ns or 'Table' in ns):
                break
            if re.match(r'^\[\d+\]', ns):
                break
            para_lines.append(ns)
            i += 1
        add_body(doc, ' '.join(para_lines))


def _render_refs(doc, block):
    lines = [l.strip() for l in block.split('\n') if l.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        if line == '---':
            i += 1
            continue
        if re.match(r'^\[\d+\]', line):
            add_ref(doc, line)
            if i + 1 < len(lines):
                nxt = lines[i + 1]
                if (nxt and not nxt.startswith('[') and not nxt.startswith('#')
                        and re.match(r'^[A-Z]', nxt)):
                    i += 1
                    add_ref(doc, lines[i])
        i += 1


if __name__ == '__main__':
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md = os.path.join(root, '论文全文.md')
    out = os.path.join(root, 'paper', '论文全文.docx')
    os.makedirs(os.path.dirname(out), exist_ok=True)
    generate_docx(md, out)
