#!/usr/bin/env python3
"""
Generate formatted DOCX for ECB-YOLO paper based on 《矿业研究与开发》 template.
Format requirements (from template PDF):
- 全文: 5号(10.5pt) 宋体, 1.5倍行距, 首行缩进2字符
- 标题: 小二号(18pt) 黑体 bold centered
- 一级标题(1,2,3...): 四号(14pt) 黑体 bold
- 二级标题(1.1,2.1...): 小四(12pt) 黑体 bold
- 三级标题(1.1.1...): 五号(10.5pt) 黑体 bold
- 摘要正文: 小五(9pt) 楷体
- 关键词: 小五(9pt) 楷体
- 参考文献: 六号(7.5pt) 宋体
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

OUTPUT = r"e:\Project\Mining_Helmet_Detection\基于ECB-YOLO的井下安全头盔检测模型.docx"
MARKDOWN = r"e:\Project\Mining_Helmet_Detection\初稿.md"


# ============================================================
# Style helpers
# ============================================================

def set_run_font(run, latin, cjk, size_pt, bold=False):
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def make_para(doc, segments, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
              line_spacing=1.5, before=0, after=0, first_line_indent=None):
    """Create a paragraph with multiple formatted runs.
    segments: list of (text, latin, cjk, size_pt, bold)
    """
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = Pt(first_line_indent)

    for seg in segments:
        text, latin, cjk, sz, bold = seg[0], seg[1], seg[2], seg[3], seg[4] if len(seg) > 4 else False
        run = para.add_run(text)
        set_run_font(run, latin, cjk, sz, bold)
    return para


# ============================================================
# Document element creators
# ============================================================

def add_title(doc, text):
    """Title: 小二号(18pt) 黑体 bold centered"""
    return make_para(doc, [(text, 'Times New Roman', '黑体', 18, True)],
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                     before=6, after=6)


def add_abstract(doc, text):
    """Abstract: label 黑体5号 bold, text 楷体9pt"""
    return make_para(doc, [
        ('摘要：', 'Times New Roman', '黑体', 10.5, True),
        (text, 'Times New Roman', '楷体', 9, False)
    ], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, before=0, after=0)


def add_keywords(doc, text):
    """Keywords: label 黑体5号 bold, text 楷体9pt"""
    return make_para(doc, [
        ('关键词：', 'Times New Roman', '黑体', 10.5, True),
        (text, 'Times New Roman', '楷体', 9, False)
    ], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, before=0, after=3)


def add_heading(doc, text, level=1):
    """Heading with level-appropriate formatting."""
    sizes = {1: 14, 2: 12, 3: 10.5}
    sb = {1: 6, 2: 4, 3: 2}
    sa = {1: 3, 2: 2, 3: 1}
    sz = sizes.get(level, 10.5)
    return make_para(doc, [(text, 'Times New Roman', '黑体', sz, True)],
                     alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                     before=sb[level], after=sa[level])


def add_section_title(doc, text):
    """Centered section title like 参考文献: 四号黑体"""
    return make_para(doc, [(text, 'Times New Roman', '黑体', 14, True)],
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                     before=6, after=6)


def add_body(doc, text, indent=True):
    """Body paragraph: 五号(10.5pt) 宋体, justified, 1.5 spacing, 2-char first-line indent."""
    fi = 2 * 10.5 if indent else None
    return make_para(doc, [(text, 'Times New Roman', '宋体', 10.5, False)],
                     alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                     before=0, after=0, first_line_indent=fi)


# ============================================================
# Page setup
# ============================================================

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Set Normal style defaults
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = 'Times New Roman'
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    # 1.5 line spacing
    ppr = style.element.get_or_add_pPr()
    spacing = ppr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:line="360" w:lineRule="auto" />')
        ppr.append(spacing)
    else:
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')


# ============================================================
# Simplified markdown parser
# ============================================================

def parse_md(filepath):
    """Parse markdown into flat list of items.
    Each item: {'type': 'title'|'h1'|'h2'|'h3'|'body'|'abstract'|'keywords', 'text': str}
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()

    items = []
    line_buffer = []

    def flush_body():
        nonlocal line_buffer
        content = '\n'.join(line_buffer).strip()
        if content:
            items.append({'type': 'body', 'text': content})
        line_buffer = []

    title_done = False
    for line in text.split('\n'):
        stripped = line.strip()

        # Title (first # line)
        if not title_done and stripped.startswith('# '):
            items.append({'type': 'title', 'text': stripped[2:].strip()})
            title_done = True
            continue

        # Headings
        if stripped.startswith('## '):
            flush_body()
            h = stripped[3:].strip()
            if h == '摘要':
                items.append({'type': 'abstract'})
            elif h.startswith('关键词') or h.startswith('关键词'):
                items.append({'type': 'keywords_header'})
            else:
                items.append({'type': 'h1', 'text': h})
            continue

        if stripped.startswith('### '):
            flush_body()
            items.append({'type': 'h2', 'text': stripped[4:].strip()})
            continue

        if stripped.startswith('#### '):
            flush_body()
            items.append({'type': 'h3', 'text': stripped[5:].strip()})
            continue

        # Skip separators
        if stripped.startswith('---'):
            flush_body()
            continue

        # Empty line = paragraph break
        if stripped == '':
            flush_body()
            continue

        line_buffer.append(stripped)

    flush_body()
    return items


# ============================================================
# Generate document
# ============================================================

def generate():
    items = parse_md(MARKDOWN)
    doc = Document()
    setup_page(doc)

    # Collect abstract text and keywords
    abstract_parts = []
    keywords_text = ''
    in_abstract = False
    in_keywords = False

    for item in items:
        if item['type'] == 'abstract':
            in_abstract = True
            in_keywords = False
        elif item['type'] == 'keywords_header':
            in_keywords = True
            in_abstract = False
        elif item['type'] == 'h1':
            break  # Stop collecting abstract when body begins
        elif item['type'] == 'body':
            if in_abstract:
                txt = item['text'].replace('**', '')
                # Check if this body contains keywords embedded
                kw_match = re.search(r'(\*\*)?关键词(\*\*)?[：:]\s*(.+)', txt)
                if kw_match:
                    before_kw = txt[:kw_match.start()].strip()
                    if before_kw:
                        abstract_parts.append(before_kw)
                    keywords_text = kw_match.group(3).strip()
                else:
                    abstract_parts.append(txt)
            elif in_keywords:
                txt = item['text'].replace('**', '')
                if not keywords_text:
                    for prefix in ['关键词：', '关键词:', 'Keywords：', 'Keywords:']:
                        if txt.startswith(prefix):
                            txt = txt[len(prefix):]
                            break
                    keywords_text = txt.strip()
                else:
                    keywords_text += ' ' + txt.strip()

    abstract_text = ''.join(abstract_parts)

    # ---- Title ----
    title_item = next((i for i in items if i['type'] == 'title'), None)
    if title_item:
        add_title(doc, title_item['text'])

    # ---- Abstract ----
    if abstract_text:
        add_abstract(doc, abstract_text)

    # ---- Keywords ----
    if keywords_text:
        add_keywords(doc, keywords_text)

    # ---- Body content ----
    in_body = False
    for item in items:
        t = item['type']

        if t == 'title':
            continue
        if t in ('abstract', 'keywords_header'):
            continue
        if t == 'body' and not in_body:
            # Before first h1, these were abstract/keywords-related
            continue

        if t == 'h1':
            in_body = True
            add_heading(doc, item['text'], level=1)
        elif t == 'h2':
            add_heading(doc, item['text'], level=2)
        elif t == 'h3':
            add_heading(doc, item['text'], level=3)
        elif t == 'body':
            text = item['text'].replace('**', '')
            add_body(doc, text)

    # ---- References section ----
    add_section_title(doc, '参考文献')

    # ---- Save ----
    doc.save(OUTPUT)
    print(f'Document saved: {OUTPUT}')
    print(f'  Title: {title_item["text"] if title_item else "N/A"}')
    print(f'  Abstract: {len(abstract_text)} chars')
    print(f'  Keywords: {keywords_text}')
    body_count = sum(1 for i in items if i['type'] == 'h1')
    print(f'  Body h1 sections: {body_count}')
    total_items = len(items)
    print(f'  Total items: {total_items}')


if __name__ == '__main__':
    generate()
